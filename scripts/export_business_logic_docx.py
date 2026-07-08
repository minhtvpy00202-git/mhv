from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = ROOT / "docs" / "TONG_HOP_CHUC_NANG_VA_LOGIC_NGHIEP_VU.md"
OUTPUT_PATH = ROOT / "docs" / "TONG_HOP_CHUC_NANG_VA_LOGIC_NGHIEP_VU.docx"


def set_a4_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


def set_default_font(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_inline_runs(paragraph, text: str) -> None:
    parts = text.split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if index % 2 == 1:
            run.bold = True


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"Heading {level}"]
    paragraph.paragraph_format.space_after = Pt(6)
    add_inline_runs(paragraph, text)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline_runs(paragraph, text)


def add_paragraph(document: Document, text: str, bold: bool = False, center: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold:
        run = paragraph.add_run(text)
        run.bold = True
    else:
        add_inline_runs(paragraph, text)


def build_document(source_text: str) -> Document:
    document = Document()
    set_a4_layout(document)
    set_default_font(document)
    add_page_number(document.sections[0])

    first_heading = True
    for raw_line in source_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            add_paragraph(document, stripped[2:].strip(), center=True)
            title_run = document.paragraphs[-1].runs[0]
            title_run.bold = True
            title_run.font.size = Pt(16)
            continue

        if stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), 1)
            if first_heading:
                first_heading = False
            continue

        if stripped.startswith("### "):
            add_heading(document, stripped[4:].strip(), 2)
            continue

        if stripped.startswith("- "):
            add_bullet(document, stripped[2:].strip())
            continue

        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            add_paragraph(document, stripped[2:-2].strip(), bold=True)
            continue

        add_paragraph(document, stripped)

    return document


def main() -> None:
    source_text = SOURCE_PATH.read_text(encoding="utf-8")
    document = build_document(source_text)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
