from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path("/Users/tranminh/FPOLY/AI/mhv/docs/PHAN_CONG_NHOM_THEO_TINH_NANG_CHI_TIET.docx")


ROWS = [
    {
        "stt": "1",
        "ten": "Minh",
        "tinh_nang_chinh": "Sơ đồ định vị tài sản",
        "tinh_nang_con": [
            "Quản lý tầng sơ đồ, tạo mới, sửa, xóa tầng.",
            "Vận hành song song 2 chế độ GRID và IMAGE, không phá logic cũ.",
            "Import ảnh làm nền cho tầng IMAGE.",
            "Vẽ phòng trên ảnh bằng rectangle và polygon, vẽ lại phòng, sửa thông tin phòng trên sơ đồ.",
            "Hiển thị marker tài sản trên sơ đồ và tìm kiếm tài sản theo mã, tên, loại, phòng.",
            "Đồng bộ dữ liệu phòng trên sơ đồ với Location, RoomShape, Asset.location.",
        ],
        "backend": [
            "src/main/java/com/poly/mhv/controller/AssetMapController.java",
            "src/main/java/com/poly/mhv/service/AssetMapService.java",
            "src/main/java/com/poly/mhv/service/AssetMapImportService.java",
            "src/main/java/com/poly/mhv/entity/MapFloor.java",
            "src/main/java/com/poly/mhv/entity/MapFloorMode.java",
            "src/main/java/com/poly/mhv/entity/RoomShape.java",
            "src/main/java/com/poly/mhv/dto/assetmap/MapFloorCreateRequest.java",
            "src/main/java/com/poly/mhv/dto/assetmap/MapFloorUpdateRequest.java",
            "src/main/java/com/poly/mhv/dto/assetmap/MapFloorResponse.java",
            "src/main/java/com/poly/mhv/dto/assetmap/RoomShapeSaveRequest.java",
            "src/main/java/com/poly/mhv/dto/assetmap/RoomShapeResponse.java",
            "src/main/java/com/poly/mhv/dto/assetmap/AssetMapPointDto.java",
            "src/main/java/com/poly/mhv/dto/assetmap/AssetMapBoundsDto.java",
            "src/main/java/com/poly/mhv/dto/assetmapimport/AssetMapImportAnalyzeResponse.java",
            "src/main/java/com/poly/mhv/dto/assetmapimport/AssetMapImportApplyRequest.java",
            "src/main/java/com/poly/mhv/dto/assetmapimport/AssetMapImportApplyResponse.java",
            "src/main/java/com/poly/mhv/dto/assetmapimport/AssetMapImportDrawingResponse.java",
            "src/main/java/com/poly/mhv/config/StaticResourceConfig.java",
        ],
        "frontend": [
            "frontend/src/pages/admin/AssetMapManagement.jsx",
            "frontend/src/components/ui/ConfirmDialog.jsx",
            "frontend/src/utils/mediaUrl.js",
            "frontend/src/App.jsx",
            "frontend/src/layouts/AdminLayout.jsx",
        ],
        "test": [
            "Tạo tầng GRID, tạo tầng IMAGE, sửa tên tầng, xóa tầng.",
            "Import ảnh nền và kiểm tra ảnh hiện lại sau khi reload trang.",
            "Thêm phòng bằng rectangle, polygon, vẽ lại phòng và lưu không refresh trang.",
            "Kiểm tra tìm kiếm tài sản và marker xuất hiện đúng phòng trên sơ đồ.",
            "Kiểm tra xóa phòng, sửa thông tin phòng, room context menu, leave guard khi thao tác đang dở dang.",
        ],
    },
    {
        "stt": "2",
        "ten": "Bạn 2",
        "tinh_nang_chinh": "Phòng - khu vực - tầng - hasAsset",
        "tinh_nang_con": [
            "CRUD phòng, khu vực trong chức năng Quản lý phòng - khu vực.",
            "Gán phòng vào tầng và đồng bộ với module sơ đồ.",
            "Quản lý cờ hasAsset true/false cho khu vực có lưu trữ tài sản hay không.",
            "Ẩn các khu vực hasAsset=false khỏi các danh sách nghiệp vụ cần phòng lưu tài sản.",
            "Chọn nhiều dòng trong bảng và xóa hàng loạt dữ liệu đã chọn.",
            "Đảm bảo ràng buộc không đổi hasAsset nếu khu vực đã có tài sản.",
        ],
        "backend": [
            "src/main/java/com/poly/mhv/controller/LocationController.java",
            "src/main/java/com/poly/mhv/service/LocationService.java",
            "src/main/java/com/poly/mhv/repository/LocationRepository.java",
            "src/main/java/com/poly/mhv/entity/Location.java",
            "src/main/java/com/poly/mhv/dto/location/LocationResponse.java",
            "src/main/java/com/poly/mhv/dto/location/LocationBulkDeleteRequest.java",
            "src/main/java/com/poly/mhv/dto/location/LocationCreateRequest.java",
            "src/main/java/com/poly/mhv/dto/location/LocationUpdateRequest.java",
            "src/main/java/com/poly/mhv/entity/MapFloor.java",
        ],
        "frontend": [
            "frontend/src/pages/admin/LocationManagement.jsx",
            "frontend/src/components/ui/ConfirmDialog.jsx",
            "frontend/src/components/ui/ActionIconButton.jsx",
            "frontend/src/hooks/useTableSort.js",
        ],
        "test": [
            "Tạo mới phòng và khu vực không chứa tài sản.",
            "Sửa hasAsset từ true sang false và ngược lại với cả 2 trường hợp có tài sản và chưa có tài sản.",
            "Kiểm tra bảng danh sách có ẩn đúng các khu vực hasAsset=false theo yêu cầu.",
            "Chọn nhiều dòng và xóa hàng loạt, kiểm tra thông báo lỗi nếu có bản ghi không hợp lệ.",
            "Kiểm tra liên kết với sơ đồ: phòng được tạo/sửa có thể được vẽ và hiện đúng.",
        ],
    },
    {
        "stt": "3",
        "ten": "Bạn 3",
        "tinh_nang_chinh": "Quản lý tài sản và lịch sử sử dụng",
        "tinh_nang_con": [
            "Tạo mới, cập nhật, xem chi tiết tài sản.",
            "Gán tài sản vào phòng hiện tại và phòng gốc.",
            "Chỉ cho phép chọn location hợp lệ có hasAsset=true trong các form liên quan.",
            "Cập nhật vị trí tài sản khi mượn, trả, điều chuyển hoặc thay đổi phòng.",
            "Theo dõi lịch sử sử dụng, lịch sử di chuyển, usage history.",
            "Phối hợp với module sơ đồ để marker tài sản hiện đúng vị trí.",
        ],
        "backend": [
            "src/main/java/com/poly/mhv/controller/AssetController.java",
            "src/main/java/com/poly/mhv/service/AssetService.java",
            "src/main/java/com/poly/mhv/repository/AssetRepository.java",
            "src/main/java/com/poly/mhv/entity/Asset.java",
            "src/main/java/com/poly/mhv/controller/UsageHistoryController.java",
            "src/main/java/com/poly/mhv/service/UsageHistoryService.java",
            "src/main/java/com/poly/mhv/repository/UsageHistoryRepository.java",
            "src/main/java/com/poly/mhv/entity/UsageHistory.java",
            "các DTO/request/response liên quan Asset và UsageHistory",
        ],
        "frontend": [
            "frontend/src/pages/admin/AssetManagement.jsx",
            "frontend/src/pages/admin/UsageHistoryManagement.jsx",
        ],
        "test": [
            "Tạo tài sản mới, kiểm tra dropdown phòng chỉ hiện khu vực hợp lệ.",
            "Sửa phòng hiện tại và phòng gốc của tài sản, kiểm tra dữ liệu lưu đúng.",
            "Mượn/trả/điều chuyển tài sản và kiểm tra lịch sử usage history.",
            "Kiểm tra tài sản hiện đúng trên sơ đồ sau khi thay đổi phòng.",
            "Kiểm tra tìm kiếm, lọc, xem chi tiết tài sản không bị ảnh hưởng bởi thay đổi location.",
        ],
    },
    {
        "stt": "4",
        "ten": "Bạn 4",
        "tinh_nang_chinh": "Kiểm kê, bảo trì, danh mục, nhà cung cấp",
        "tinh_nang_con": [
            "Quản lý đợt kiểm kê, phiên kiểm kê và kết quả kiểm kê.",
            "Đảm bảo các luồng kiểm kê chỉ sử dụng phòng hợp lệ nếu cần chọn phòng.",
            "Quản lý lịch sử bảo trì, sửa chữa, thông tin kỹ thuật của tài sản.",
            "Quản lý danh mục loại tài sản.",
            "Quản lý nhà cung cấp và thông tin liên quan.",
            "Phối hợp dữ liệu với module tài sản để bảo trì/kiểm kê đọc đúng dữ liệu thực tế.",
        ],
        "backend": [
            "src/main/java/com/poly/mhv/controller/InventoryAuditController.java",
            "src/main/java/com/poly/mhv/service/InventoryAuditService.java",
            "src/main/java/com/poly/mhv/repository/InventoryAuditRepository.java",
            "src/main/java/com/poly/mhv/entity/InventoryAudit.java",
            "src/main/java/com/poly/mhv/controller/MaintenanceHistoryController.java",
            "src/main/java/com/poly/mhv/service/MaintenanceHistoryService.java",
            "src/main/java/com/poly/mhv/repository/MaintenanceHistoryRepository.java",
            "src/main/java/com/poly/mhv/entity/MaintenanceHistory.java",
            "src/main/java/com/poly/mhv/controller/CategoryController.java",
            "src/main/java/com/poly/mhv/service/CategoryService.java",
            "src/main/java/com/poly/mhv/repository/CategoryRepository.java",
            "src/main/java/com/poly/mhv/entity/Category.java",
            "src/main/java/com/poly/mhv/controller/SupplierController.java",
            "src/main/java/com/poly/mhv/service/SupplierService.java",
            "src/main/java/com/poly/mhv/repository/SupplierRepository.java",
            "src/main/java/com/poly/mhv/entity/Supplier.java",
        ],
        "frontend": [
            "frontend/src/pages/admin/InventoryAuditManagement.jsx",
            "frontend/src/pages/admin/MaintenanceHistoryManagement.jsx",
            "frontend/src/pages/admin/CategoryManagement.jsx",
            "frontend/src/pages/admin/SupplierManagement.jsx",
        ],
        "test": [
            "Tạo đợt kiểm kê, lập phiên kiểm kê và chạy luồng cập nhật kết quả.",
            "Kiểm tra các dropdown chọn tài sản, phòng, danh mục vẫn hoạt động đúng.",
            "Thêm/sửa/xóa lịch sử bảo trì và kiểm tra liên kết với tài sản.",
            "Thêm/sửa/xóa danh mục và nhà cung cấp, kiểm tra dữ liệu hiển thị ở form tài sản.",
            "Kiểm tra không có chức năng nào bị ảnh hưởng sau khi lọc khu vực hasAsset=false.",
        ],
    },
    {
        "stt": "5",
        "ten": "Bạn 5",
        "tinh_nang_chinh": "Người dùng, phân quyền, hỗ trợ kỹ thuật, branding",
        "tinh_nang_con": [
            "Quản lý tài khoản người dùng và vai trò.",
            "Phân quyền truy cập theo chức năng và trang.",
            "Quản lý loại hỗ trợ kỹ thuật và các danh mục liên quan.",
            "Quản lý cấu hình thương hiệu: tên hệ thống, tên công ty, màu sắc, thông tin hiển thị.",
            "Theo dõi giao diện đăng nhập, admin layout, các route có bảo vệ xác thực.",
            "Đảm bảo route, auth, security và branding API hoạt động ổn định.",
        ],
        "backend": [
            "src/main/java/com/poly/mhv/controller/UserController.java",
            "src/main/java/com/poly/mhv/service/UserService.java",
            "src/main/java/com/poly/mhv/repository/UserRepository.java",
            "src/main/java/com/poly/mhv/entity/User.java",
            "src/main/java/com/poly/mhv/entity/Role.java",
            "src/main/java/com/poly/mhv/controller/TechSupportTypeController.java",
            "src/main/java/com/poly/mhv/service/TechSupportTypeService.java",
            "src/main/java/com/poly/mhv/repository/TechSupportTypeRepository.java",
            "src/main/java/com/poly/mhv/entity/TechSupportType.java",
            "src/main/java/com/poly/mhv/controller/BrandingSettingsController.java",
            "src/main/java/com/poly/mhv/service/BrandingService.java",
            "src/main/java/com/poly/mhv/entity/AppSettings.java",
            "src/main/java/com/poly/mhv/repository/AppSettingsRepository.java",
            "src/main/java/com/poly/mhv/config/SecurityConfig.java",
        ],
        "frontend": [
            "frontend/src/pages/admin/UserManagement.jsx",
            "frontend/src/pages/admin/TechSupportTypeManagement.jsx",
            "frontend/src/contexts/BrandingContext.jsx",
            "frontend/src/contexts/AuthContext.jsx",
            "frontend/src/layouts/AdminLayout.jsx",
            "frontend/src/pages/auth/LoginPage.jsx",
        ],
        "test": [
            "Đăng nhập/đăng xuất với các vai trò khác nhau.",
            "Kiểm tra các trang admin có ẩn/hiện đúng theo quyền.",
            "Thêm/sửa/xóa người dùng và loại hỗ trợ kỹ thuật.",
            "Cập nhật branding và kiểm tra login, sidebar, title, màu sắc thay đổi đúng.",
            "Kiểm tra route và API branding/security hoạt động ổn định trên local và production.",
        ],
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("PHÂN CÔNG NHÓM THEO TÍNH NĂNG - BẢN SIÊU CHI TIẾT")
    run.bold = True
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run(
        "Format: Stt | Tên | Tính năng chính | Tính năng con | File backend | File frontend | Luồng test cần tự kiểm tra"
    )
    run.italic = True
    run.font.size = Pt(9)

    note = doc.add_paragraph()
    run = note.add_run(
        "Lưu ý: tên thành viên hiện để tạm là Minh, Bạn 2, Bạn 3, Bạn 4, Bạn 5 để có thể đổi nhanh sau."
    )
    run.italic = True
    run.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Stt",
        "Tên",
        "Tính năng chính",
        "Tính năng con",
        "File backend",
        "File frontend",
        "Luồng test cần tự kiểm tra",
    ]
    widths = [0.4, 0.8, 1.3, 2.8, 3.1, 2.4, 2.9]

    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for item in ROWS:
        row = table.add_row().cells
        values = [
            item["stt"],
            item["ten"],
            item["tinh_nang_chinh"],
            item["tinh_nang_con"],
            item["backend"],
            item["frontend"],
            item["test"],
        ]

        for index, value in enumerate(values):
            cell = row[index]
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            if isinstance(value, list):
                cell.text = ""
                for item_index, line in enumerate(value):
                    paragraph = cell.paragraphs[0] if item_index == 0 else cell.add_paragraph()
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    run = paragraph.add_run(f"- {line}")
                    run.font.size = Pt(8)
            else:
                cell.text = value
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    if index == 0:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    footer = doc.add_paragraph()
    run = footer.add_run(
        "Có thể đưa file này cho cả nhóm để giao việc; mỗi thành viên tự phụ trách backend, frontend và test end-to-end phạm vi của mình."
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
